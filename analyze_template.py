from docx import Document
import os

template_path = 'template.docx'
output_file = 'template_analysis.txt'

if os.path.exists(template_path):
    try:
        doc = Document(template_path)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("=== TEMPLATE.DOCX STRUCTURE ANALYSIS ===\n\n")
            
            f.write("PARAGRAPH-BY-PARAGRAPH WITH STYLES:\n")
            f.write("=" * 80 + "\n\n")
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    style = para.style.name if para.style else "No Style"
                    f.write(f"[Para {i:03d}] [Style: {style}]\n")
                    f.write(f"{text}\n")
                    f.write("-" * 80 + "\n\n")
            
            f.write("\n\n=== FULL RAW CONTENT ===\n")
            f.write("=" * 80 + "\n\n")
            
            for para in doc.paragraphs:
                if para.text.strip():
                    f.write(para.text + "\n")
        
        print(f"Analysis saved to: {output_file}")
        
        # Also print summary
        print("\n=== SUMMARY ===")
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        print(f"Non-empty paragraphs: {sum(1 for p in doc.paragraphs if p.text.strip())}")
        
    except Exception as e:
        print(f"Error: {e}")
else:
    print(f"File not found: {template_path}")
